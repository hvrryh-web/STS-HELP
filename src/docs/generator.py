"""
Documentation generator for Monte Carlo simulation results.

Generates deliverables in multiple formats:
- PDF: Print-quality report
- DOCX: Editable version with methodology notes
- XLSX: Raw data with Deck_Eval_Model sheet
"""

from dataclasses import dataclass
from typing import List, Dict, Any, Optional
from pathlib import Path
from datetime import datetime
import json

# PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors

# DOCX generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# XLSX generation
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter


class DocumentGenerator:
    """
    Generates documentation in PDF, DOCX, and XLSX formats.
    
    Attributes:
        simulation_results: Dictionary containing simulation results
        output_dir: Directory to save generated documents
    """
    
    def __init__(self, simulation_results: Dict[str, Any], output_dir: Path = None):
        """
        Initialize the documentation generator.
        
        Args:
            simulation_results: Results from Monte Carlo simulation
            output_dir: Output directory (default: ./outputs)
        """
        self.results = simulation_results
        
        if output_dir is None:
            output_dir = Path(__file__).parent.parent.parent / "outputs"
        
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    def generate_pdf(self, filename: str = None) -> Path:
        """
        Generate PDF report.
        
        Args:
            filename: Output filename (default: auto-generated)
            
        Returns:
            Path to generated PDF file
        """
        if filename is None:
            filename = f"simulation_report_{self.timestamp}.pdf"
        
        output_path = self.output_dir / filename
        
        # Create PDF document
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=1*inch,
            bottomMargin=0.75*inch
        )
        
        # Build content
        story = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f4788'),
            spaceAfter=30
        )
        
        story.append(Paragraph("Monte Carlo Simulation Report", title_style))
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", styles['Heading2']))
        story.append(Spacer(1, 0.1*inch))
        
        summary_text = self._generate_summary_text()
        story.append(Paragraph(summary_text, styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Simulation Parameters
        story.append(Paragraph("Simulation Parameters", styles['Heading2']))
        story.append(Spacer(1, 0.1*inch))
        
        params_data = self._get_parameters_table_data()
        params_table = Table(params_data, colWidths=[2.5*inch, 3*inch])
        params_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(params_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Results by Scenario
        story.append(Paragraph("Results by Scenario", styles['Heading2']))
        story.append(Spacer(1, 0.1*inch))
        
        results_data = self._get_results_table_data()
        results_table = Table(results_data, colWidths=[1.2*inch, 1.2*inch, 1.2*inch, 1.2*inch, 1.2*inch])
        results_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.lightblue),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
        ]))
        story.append(results_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Methodology
        story.append(PageBreak())
        story.append(Paragraph("Methodology", styles['Heading2']))
        story.append(Spacer(1, 0.1*inch))
        
        methodology_text = self._generate_methodology_text()
        for para in methodology_text.split('\n\n'):
            if para.strip():
                story.append(Paragraph(para.strip(), styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(story)
        
        return output_path
    
    def generate_docx(self, filename: str = None) -> Path:
        """
        Generate DOCX report.
        
        Args:
            filename: Output filename (default: auto-generated)
            
        Returns:
            Path to generated DOCX file
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX generation")
        
        if filename is None:
            filename = f"simulation_report_{self.timestamp}.docx"
        
        output_path = self.output_dir / filename
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading('Monte Carlo Simulation Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading('Executive Summary', 1)
        doc.add_paragraph(self._generate_summary_text())
        
        # Simulation Parameters
        doc.add_heading('Simulation Parameters', 1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parameter'
        hdr_cells[1].text = 'Value'
        
        params_data = self._get_parameters_table_data()
        for row_data in params_data[1:]:  # Skip header
            row_cells = table.add_row().cells
            row_cells[0].text = str(row_data[0])
            row_cells[1].text = str(row_data[1])
        
        doc.add_paragraph()
        
        # Results
        doc.add_heading('Results by Scenario', 1)
        results_table = doc.add_table(rows=1, cols=5)
        results_table.style = 'Light Grid Accent 1'
        
        results_data = self._get_results_table_data()
        hdr_cells = results_table.rows[0].cells
        for i, val in enumerate(results_data[0]):
            hdr_cells[i].text = str(val)
        
        for row_data in results_data[1:]:
            row_cells = results_table.add_row().cells
            for i, val in enumerate(row_data):
                row_cells[i].text = str(val)
        
        doc.add_paragraph()
        
        # Methodology
        doc.add_page_break()
        doc.add_heading('Methodology', 1)
        methodology_text = self._generate_methodology_text()
        for para in methodology_text.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Save document
        doc.save(str(output_path))
        
        return output_path
    
    def generate_xlsx(self, filename: str = None) -> Path:
        """
        Generate XLSX report with raw data.
        
        Args:
            filename: Output filename (default: auto-generated)
            
        Returns:
            Path to generated XLSX file
        """
        if filename is None:
            filename = f"simulation_data_{self.timestamp}.xlsx"
        
        output_path = self.output_dir / filename
        
        # Create workbook
        wb = Workbook()
        
        # Summary sheet
        ws_summary = wb.active
        ws_summary.title = "Summary"
        
        self._create_summary_sheet(ws_summary)
        
        # Results by scenario sheet
        ws_results = wb.create_sheet("Results_by_Scenario")
        self._create_results_sheet(ws_results)
        
        # Deck_Eval_Model sheet (if deck evaluations exist)
        if "deck_evaluations" in self.results:
            ws_deck = wb.create_sheet("Deck_Eval_Model")
            self._create_deck_eval_sheet(ws_deck)
        
        # Save workbook
        wb.save(str(output_path))
        
        return output_path
    
    def generate_all(self, prefix: str = None) -> Dict[str, Path]:
        """
        Generate all document formats.
        
        Args:
            prefix: Filename prefix (default: auto-generated)
            
        Returns:
            Dictionary mapping format to file path
        """
        if prefix is None:
            prefix = f"simulation_{self.timestamp}"
        
        outputs = {
            "pdf": self.generate_pdf(f"{prefix}.pdf"),
            "xlsx": self.generate_xlsx(f"{prefix}.xlsx"),
        }
        
        if DOCX_AVAILABLE:
            outputs["docx"] = self.generate_docx(f"{prefix}.docx")
        
        return outputs
    
    def _generate_summary_text(self) -> str:
        """Generate executive summary text."""
        if "scenarios" in self.results:
            # Multiple scenarios
            scenarios = list(self.results["scenarios"].keys())
            return (
                f"This report presents results from Monte Carlo simulations "
                f"across {len(scenarios)} scenarios: {', '.join(scenarios)}. "
                f"Simulations were run with deterministic PRNG seeds for reproducibility."
            )
        else:
            # Single scenario
            return (
                "This report presents results from Monte Carlo simulation "
                "with deterministic PRNG seed for reproducibility."
            )
    
    def _generate_methodology_text(self) -> str:
        """Generate methodology section text."""
        return """
Monte Carlo Simulation Framework

This simulation uses a deterministic PRNG (Pseudo-Random Number Generator) approach with explicit seed management to ensure reproducible results across multiple runs.

Scenario Types:
- BASE: Deterministic behavior with minimal heuristics, establishing a baseline for comparison
- COMPLEX: Full heuristic activation with synergy-aware decision making
- IDEAL: Upper-bound performance estimation using perfect-play oracle model
- RANDOM: Degraded decision logic for robustness stress-testing

Metrics Collected:
- Win Rate: Percentage of successful simulation runs
- Median Turns: Median number of turns taken across all runs
- Variance: Statistical variance in turn count
- Tail Risk (5%): Analysis of worst-case outcomes (bottom 5% percentile)

The simulation framework implements explicit deck/hand/discard mechanics, intent-aware defensive play, and character-specific card interactions.

Data Verification:
All card mechanics, relic interactions, and enemy AI patterns have been cross-referenced against verified community sources including SpireSpy, Tiny Helper, and STSDeckAssistant.
"""
    
    def _get_parameters_table_data(self) -> List[List[str]]:
        """Get simulation parameters as table data."""
        data = [["Parameter", "Value"]]
        
        if "scenarios" in self.results:
            # Get from first scenario
            first_scenario = next(iter(self.results["scenarios"].values()))
            data.extend([
                ["Iterations per Scenario", str(first_scenario.get("iterations", "N/A"))],
                ["Root Seed", str(first_scenario.get("root_seed", "N/A"))],
                ["Character", str(first_scenario.get("character", "N/A"))],
                ["Scenarios", ", ".join(self.results["scenarios"].keys())],
            ])
        else:
            data.extend([
                ["Iterations", str(self.results.get("iterations", "N/A"))],
                ["Root Seed", str(self.results.get("root_seed", "N/A"))],
                ["Character", str(self.results.get("character", "N/A"))],
                ["Scenario", str(self.results.get("scenario", "N/A"))],
            ])
        
        return data
    
    def _get_results_table_data(self) -> List[List[str]]:
        """Get results as table data."""
        data = [["Scenario", "Win Rate (%)", "Median Turns", "Mean Turns", "Std Dev"]]
        
        if "scenarios" in self.results:
            for scenario_name, scenario_data in self.results["scenarios"].items():
                data.append([
                    scenario_name,
                    f"{scenario_data.get('win_rate', 0):.2f}",
                    f"{scenario_data.get('median_turns', 0):.1f}",
                    f"{scenario_data.get('mean_turns', 0):.1f}",
                    f"{scenario_data.get('std_turns', 0):.2f}",
                ])
        else:
            data.append([
                self.results.get("scenario", "N/A"),
                f"{self.results.get('win_rate', 0):.2f}",
                f"{self.results.get('median_turns', 0):.1f}",
                f"{self.results.get('mean_turns', 0):.1f}",
                f"{self.results.get('std_turns', 0):.2f}",
            ])
        
        return data
    
    def _create_summary_sheet(self, ws):
        """Create summary sheet in Excel."""
        # Header
        ws['A1'] = 'Monte Carlo Simulation Summary'
        ws['A1'].font = Font(size=16, bold=True)
        ws.merge_cells('A1:D1')
        
        ws['A3'] = 'Generated:'
        ws['B3'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # Parameters
        row = 5
        ws[f'A{row}'] = 'Simulation Parameters'
        ws[f'A{row}'].font = Font(size=12, bold=True)
        
        params_data = self._get_parameters_table_data()
        for i, (param, value) in enumerate(params_data[1:], start=row+1):
            ws[f'A{i}'] = param
            ws[f'B{i}'] = value
        
        # Style
        for cell in ws['A']:
            cell.font = Font(bold=True)
    
    def _create_results_sheet(self, ws):
        """Create results sheet in Excel."""
        results_data = self._get_results_table_data()
        
        # Write data
        for row_idx, row_data in enumerate(results_data, start=1):
            for col_idx, value in enumerate(row_data, start=1):
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.value = value
                
                # Style header row
                if row_idx == 1:
                    cell.font = Font(bold=True, color='FFFFFF')
                    cell.fill = PatternFill(start_color='1F4788', end_color='1F4788', fill_type='solid')
                    cell.alignment = Alignment(horizontal='center')
        
        # Auto-adjust column widths
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column].width = adjusted_width
    
    def _create_deck_eval_sheet(self, ws):
        """Create Deck_Eval_Model sheet in Excel."""
        ws['A1'] = 'Deck Evaluation Model'
        ws['A1'].font = Font(size=14, bold=True)
        
        # Formula explanation
        ws['A3'] = 'Composite Scoring Formula:'
        ws['B3'] = 'Score = 0.4Q + 0.25S + 0.15C + 0.10K + 0.10R'
        
        ws['A5'] = 'Component'
        ws['B5'] = 'Weight'
        ws['C5'] = 'Description'
        
        components = [
            ('Q - Card Quality', '0.4', 'Average tier score of cards'),
            ('S - Synergy Coherence', '0.25', 'Synergy density between cards'),
            ('C - Curve Smoothness', '0.15', 'Mana curve distribution'),
            ('K - Consistency', '0.10', 'Deck size and redundancy'),
            ('R - Relic Impact', '0.10', 'Relic synergy with deck'),
        ]
        
        for i, (comp, weight, desc) in enumerate(components, start=6):
            ws[f'A{i}'] = comp
            ws[f'B{i}'] = weight
            ws[f'C{i}'] = desc
        
        # Style
        ws['A5'].font = Font(bold=True)
        ws['B5'].font = Font(bold=True)
        ws['C5'].font = Font(bold=True)


def generate_documentation(
    simulation_results: Dict[str, Any],
    output_dir: Path = None,
    formats: List[str] = None
) -> Dict[str, Path]:
    """
    Convenience function to generate documentation.
    
    Args:
        simulation_results: Results from simulation
        output_dir: Output directory
        formats: List of formats to generate (default: all)
        
    Returns:
        Dictionary mapping format to file path
    """
    generator = DocumentGenerator(simulation_results, output_dir)
    
    if formats is None:
        return generator.generate_all()
    
    outputs = {}
    for fmt in formats:
        if fmt.lower() == "pdf":
            outputs["pdf"] = generator.generate_pdf()
        elif fmt.lower() == "docx":
            outputs["docx"] = generator.generate_docx()
        elif fmt.lower() == "xlsx":
            outputs["xlsx"] = generator.generate_xlsx()
    
    return outputs
